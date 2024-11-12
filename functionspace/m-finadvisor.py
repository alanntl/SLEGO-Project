# Import required libraries
import os
import json
import pandas as pd
from typing import Union, Dict, Any, List
from openai import OpenAI
from bs4 import BeautifulSoup
import yfinance as yf
from yfinance import Ticker as si
import datetime



# =====================================
# Internal Utility Functions
# =====================================

def __check_module(module_name: str) -> bool:
    """Checks if a Python module is installed."""
    import importlib.util
    return importlib.util.find_spec(module_name) is not None

def __ensure_directory_exists(directory_path: str) -> None:
    """Ensures the directory exists; creates it if it doesn’t."""
    os.makedirs(directory_path, exist_ok=True)

def __extract_text_from_html(html_content: str) -> str:
    """Extracts and cleans text from HTML content by removing scripts and styles."""
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html_content, 'html.parser')
        for script in soup(["script", "style"]):
            script.decompose()  # Remove script and style elements
        return soup.get_text(separator='\n', strip=True)
    except Exception as e:
        raise ValueError(f"Error processing HTML: {str(e)}")

def __read_file_content(file_path: str, installed_modules: Dict[str, bool]) -> str:
    """Reads content from various file types; uses BeautifulSoup for HTML, pandas for CSV/Excel, pdfplumber for PDF, etc."""
    file_extension = os.path.splitext(file_path)[1].lower()
    
    try:
        if file_extension == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        
        elif file_extension == '.json':
            with open(file_path, 'r', encoding='utf-8') as f:
                return json.dumps(json.load(f), indent=2)
        
        elif file_extension == '.html':
            with open(file_path, 'r', encoding='utf-8') as f:
                return __extract_text_from_html(f.read())
        
        elif file_extension == '.pdf' and installed_modules.get('pdfplumber'):
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                return '\n'.join(page.extract_text() for page in pdf.pages if page.extract_text())
        
        elif file_extension == '.docx' and installed_modules.get('docx'):
            import docx
            doc = docx.Document(file_path)
            return '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
        
        elif file_extension in ['.csv', '.xlsx'] and installed_modules.get('pandas'):
            df = pd.read_csv(file_path) if file_extension == '.csv' else pd.read_excel(file_path)
            return f"Data Summary:\n{df.describe().to_string()}\n\nPreview:\n{df.head().to_string()}\n\nShape: {df.shape}"
            
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
            
    except Exception as e:
        raise ValueError(f"Error reading {file_extension} file: {str(e)}")


# =====================================
# OpenAI Content Analysis Functions
# =====================================

def __analyze_content(client: OpenAI, content: str, filename: str, 
                     system_role: str, analysis_task: str, 
                     additional_instructions: str, model: str,
                     max_tokens: int, temperature: float, 
                     chunk_size: int = 4000) -> str:
    """Private function to analyze content using OpenAI API."""
    try:
        content_chunks = [content[i:i + chunk_size] for i in range(0, len(content), chunk_size)]
        full_analysis = []

        for i, chunk in enumerate(content_chunks):
            messages = [
                {"role": "system", "content": system_role},
                {"role": "user", "content": f"""
Analysis Task: {analysis_task}
Additional Instructions: {additional_instructions}
File: {filename} (Part {i+1}/{len(content_chunks)})

Content:
{chunk}
                """}
            ]
            
            response = client.chat.completions.create(
                model=model,
                messages=messages,
                max_tokens=max_tokens,
                temperature=temperature
            )
            
            full_analysis.append(response.choices[0].message.content)
        
        if len(full_analysis) > 1:
            consolidation_messages = [
                {"role": "system", "content": system_role},
                {"role": "user", "content": f"""
Please provide a consolidated analysis summary focusing on:
1. Key Points and Insights
2. Overall Patterns and Trends
3. Notable Findings
4. Recommendations
5. Additional Observations

Previous analyses:
{' '.join(full_analysis)}
                """}
            ]
            
            summary_response = client.chat.completions.create(
                model=model,
                messages=consolidation_messages,
                max_tokens=max_tokens,
                temperature=temperature
            )
            
            return summary_response.choices[0].message.content
        
        return full_analysis[0]
        
    except Exception as e:
        return f"Error during analysis: {str(e)}"


def universal_analyzer(
    input_paths: Union[str, List[str]] = 'dataspace/input/',
    output_file: str = 'dataspace/analysis_results.json',
    api_key: str = '',
    system_role: str = "You are a highly skilled analyst with expertise in multiple domains.",
    analysis_task: str = "Analyze this content and provide key insights, patterns, and recommendations.",
    additional_instructions: str = "",
    model: str = "gpt-4o",
    max_tokens: int = 2000,
    temperature: float = 0.7,
    consolidated_analysis: bool = True
) -> Dict:
    """
    Universal content analyzer supporting multiple input files and directories.
    
    Parameters:
    - input_paths: String path or list of paths to analyze
    - output_file: Where to save analysis results
    - api_key: OpenAI API key
    - system_role: Role description for the AI
    - analysis_task: Main analysis task description
    - additional_instructions: Extra analysis instructions
    - model: OpenAI model to use
    - max_tokens: Maximum tokens in response
    - temperature: Response randomness (0-1)
    - consolidated_analysis: Whether to provide a consolidated analysis of all files
    
    Returns:
    - Dict containing analysis results and metadata
    """
    try:
        # Initialize modules and file types
        installed_modules = {
            module: __check_module(module)
            for module in ['docx', 'pdfplumber', 'pandas', 'openpyxl', 'bs4']
        }

        supported_file_types = ['.txt', '.json', '.html']
        if installed_modules.get('docx'):
            supported_file_types.append('.docx')
        if installed_modules.get('pdfplumber'):
            supported_file_types.append('.pdf')
        if installed_modules.get('pandas'):
            supported_file_types.extend(['.csv', '.xlsx'])

        if not api_key:
            raise ValueError("OpenAI API key is required")

        client = OpenAI(api_key=api_key)
        analysis_results = {
            'system_info': {
                'installed_modules': installed_modules,
                'supported_file_types': supported_file_types
            },
            'files': {},
            'consolidated_analysis': None
        }

        def analyze_file(file_path: str) -> Dict:
            """Analyze a single file."""
            file_extension = os.path.splitext(file_path)[1].lower()
            if file_extension in supported_file_types:
                content = __read_file_content(file_path, installed_modules)
                analysis = __analyze_content(
                    client=client,
                    content=content,
                    filename=os.path.basename(file_path),
                    system_role=system_role,
                    analysis_task=analysis_task,
                    additional_instructions=additional_instructions,
                    model=model,
                    max_tokens=max_tokens,
                    temperature=temperature
                )
                return {
                    'file_path': file_path,
                    'analysis': analysis,
                    'file_type': file_extension
                }
            return {
                'file_path': file_path,
                'error': f"Unsupported file type: {file_extension}",
                'file_type': file_extension
            }

        def process_path(path: str) -> List[Dict]:
            """Process a single path (file or directory)."""
            results = []
            if os.path.isfile(path):
                results.append(analyze_file(path))
            elif os.path.isdir(path):
                for root, _, files in os.walk(path):
                    for file in files:
                        file_path = os.path.join(root, file)
                        results.append(analyze_file(file_path))
            return results

        # Process all input paths
        all_results = []
        input_paths = [input_paths] if isinstance(input_paths, str) else input_paths
        
        for path in input_paths:
            path_results = process_path(path)
            all_results.extend(path_results)
            
        # Store individual file analyses
        for result in all_results:
            file_name = os.path.basename(result['file_path'])
            analysis_results['files'][file_name] = result

        # Generate consolidated analysis if requested and there are multiple files
        if consolidated_analysis and len(all_results) > 1:
            successful_analyses = [r for r in all_results if 'error' not in r]
            if successful_analyses:
                combined_content = "\n\n".join([
                    f"File: {os.path.basename(r['file_path'])} (Type: {r['file_type']})\n"
                    f"Analysis:\n{r['analysis']}"
                    for r in successful_analyses
                ])
                
                consolidated = __analyze_content(
                    client=client,
                    content=combined_content,
                    filename="All Files",
                    system_role=system_role,
                    analysis_task=f"{analysis_task}\n\nProvide a consolidated analysis of all files.",
                    additional_instructions=f"{additional_instructions}\n\nFocus on connections and patterns across all files.",
                    model=model,
                    max_tokens=max_tokens,
                    temperature=temperature
                )
                
                analysis_results['consolidated_analysis'] = consolidated

        # Save results
        __ensure_directory_exists(output_file)
        if os.path.exists(output_file):
            os.remove(output_file)  # Delete existing file if it exists
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(analysis_results, f, indent=4, ensure_ascii=False)
            
        return analysis_results

    except Exception as e:
        error_result = {'error': str(e)}
        __ensure_directory_exists(output_file)
        if os.path.exists(output_file):
            os.remove(output_file)  # Delete existing file if it exists
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(error_result, f, indent=4, ensure_ascii=False)
        return error_result



def financial_advisor(
    input_paths: Union[str, List[str]] = 'dataspace/input/',
    output_file: str = 'dataspace/financial_advice.json',
    api_key: str = '',
    investment_context: str = "General investment analysis and recommendations",
    risk_profile: str = "Moderate",
    time_horizon: str = "Long-term (5+ years)",
    analysis_task: str = "Provide comprehensive investment analysis and recommendations",
    model: str = "gpt-4o",
    max_tokens: int = 2000,
    temperature: float = 0.7,
    consolidated_analysis: bool = True
) -> Dict:
    """
    Analyzes multiple financial documents and provides professional investment advice.
    
    Parameters:
    - input_paths: String path or list of paths to financial documents
    - output_file: Where to save the analysis and recommendations
    - api_key: OpenAI API key
    - investment_context: Specific context or goals for the investment analysis
    - risk_profile: Investment risk profile (e.g., Conservative, Moderate, Aggressive)
    - time_horizon: Investment time horizon
    - analysis_task: Specific analysis task or focus
    - model: OpenAI model to use
    - max_tokens: Maximum tokens in response
    - temperature: Response creativity (0-1)
    - consolidated_analysis: Whether to provide a consolidated analysis of all documents
    """
    return universal_analyzer(  input_paths=input_paths,
                                output_file=output_file,
                                api_key=api_key,
                                system_role="""You are a highly experienced financial advisor with expertise in:
                                - Investment analysis and portfolio management
                                - Market analysis and trend identification
                                - Risk assessment and management
                                - Financial planning and strategy development""",
                                analysis_task=f"""Investment Context: {investment_context}
                        Risk Profile: {risk_profile}
                        Time Horizon: {time_horizon}

                        {analysis_task}""",
                                additional_instructions="""Provide thorough, professional analysis considering:
                        - The client's risk profile and investment goals
                        - Current market conditions and trends
                        - Potential risks and mitigation strategies
                        - Long-term investment implications""",
                                model=model,
                                max_tokens=max_tokens,
                                temperature=temperature,
                                consolidated_analysis=consolidated_analysis
                            )


# =====================================
# YFinance Data Functions
# =====================================

import pandas as pd
import yfinance as yf
import os
from datetime import datetime

def get_all_financial_data(
    ticker_symbol: str = 'TSLA',
    period_for_market_data: str = '1y',
    output_folder: str = 'dataspace/financials',
    save_as_excel: bool = False
) -> dict:
    """
    Fetches comprehensive financial data for a given stock ticker and saves it to files.
    
    Parameters:
    -----------
    ticker_symbol : str
        The stock ticker symbol (e.g., 'MSFT' for Microsoft)
    period_for_market_data : str
        The period for historical data (e.g., '1d', '5d', '1mo', '3mo', '6mo', '1y', '2y', '5y', '10y', 'ytd', 'max')
    output_folder : str
        The folder where output files will be saved
    save_as_excel : bool
        If True, saves data to Excel file; if False, saves to separate CSV files
    
    Returns:
    --------
    dict
        Dictionary containing all the collected financial data
    """
    
    # Create the output folder if it doesn't exist
    os.makedirs(output_folder, exist_ok=True)
    
    # Initialize the Ticker object
    ticker = yf.Ticker(ticker_symbol)
    
    # Dictionary to hold all data
    data = {}
    
    # Helper function to convert DataFrame to timezone-naive
    def make_timezone_naive(df):
        if not isinstance(df, pd.DataFrame):
            if isinstance(df, pd.Series):
                df = df.to_frame()
            else:
                return df
                
        # Convert index if it's datetime with timezone
        if isinstance(df.index, pd.DatetimeIndex) and df.index.tz is not None:
            df.index = df.index.tz_localize(None)
        
        # Convert columns with timezones
        for col in df.select_dtypes(include=['datetime64[ns, UTC]', 'datetimetz']).columns:
            df[col] = df[col].dt.tz_localize(None)
            
        # Convert any datetime objects in object columns
        for col in df.select_dtypes(include=['object']).columns:
            try:
                if pd.api.types.is_datetime64_any_dtype(df[col]):
                    df[col] = pd.to_datetime(df[col]).dt.tz_localize(None)
            except:
                continue
                
        return df

    # Helper function to safely fetch and process data
    def safe_fetch(data_name, fetch_func):
        try:
            result = fetch_func()
            if isinstance(result, (pd.DataFrame, pd.Series)):
                return make_timezone_naive(result)
            elif isinstance(result, dict):
                df = pd.DataFrame.from_dict(result, orient='index', columns=['Value'])
                return make_timezone_naive(df)
            elif result is not None:
                df = pd.DataFrame([result])
                return make_timezone_naive(df)
            return pd.DataFrame()
        except Exception as e:
            print(f"Warning: Failed to fetch {data_name}: {str(e)}")
            return pd.DataFrame()

    # Fetch and store various data
    data_fetchers = {
        'info': lambda: pd.DataFrame.from_dict(ticker.info, orient='index', columns=[ticker_symbol]),
        'history': lambda: ticker.history(period=period_for_market_data),
        'news': lambda: pd.DataFrame(ticker.news),
        'actions': lambda: ticker.actions if not isinstance(ticker.actions, pd.Series) else ticker.actions.to_frame(),
        'dividends': lambda: ticker.dividends.to_frame() if not ticker.dividends.empty else pd.DataFrame(),
        'splits': lambda: ticker.splits.to_frame() if not ticker.splits.empty else pd.DataFrame(),
        'capital_gains': lambda: ticker.capital_gains.to_frame() if not ticker.capital_gains.empty else pd.DataFrame(),
        #pd.DataFrame(ticker.calendar)
        'calendar': lambda: pd.DataFrame(ticker.calendar),
        # sec_filings
        'sec_filings': lambda: pd.DataFrame(ticker.sec_filings),



        'income_statement': lambda: ticker.income_stmt.transpose(),
        'quarterly_income_statement': lambda: ticker.quarterly_income_stmt.transpose(),
        'balance_sheet': lambda: ticker.balance_sheet.transpose(),
        'quarterly_balance_sheet': lambda: ticker.quarterly_balance_sheet.transpose(),
        'cashflow': lambda: ticker.cashflow.transpose(),
        'quarterly_cashflow': lambda: ticker.quarterly_cashflow.transpose(),


        'major_holders': lambda: ticker.major_holders,
        'institutional_holders': lambda: ticker.institutional_holders,
        'mutualfund_holders': lambda: ticker.mutualfund_holders,
         #msft.insider_transactions
        'insider_transactions': lambda: ticker.insider_transactions,
        #msft.insider_purchases
        'insider_purchases': lambda: ticker.insider_purchases,
        #msft.insider_roster_holders
        'insider_roster_holders': lambda: ticker.insider_roster_holders,

        'sustainability': lambda: ticker.sustainability.transpose(),
        'recommendations': lambda: ticker.recommendations,
        'recommendations_summary': lambda: ticker.recommendations_summary,
        #msft.upgrades_downgrades
        'upgrades_downgrades': lambda: ticker.upgrades_downgrades,
        'analyst_price_targets': lambda: ticker.analyst_price_targets,
        'earnings_estimate': lambda: ticker.earnings_estimate,
        'revenue_estimate': lambda: ticker.revenue_estimate,
        'earnings_history': lambda: ticker.earnings_history,
        'eps_trend': lambda: ticker.eps_trend,
        'eps_revisions': lambda: ticker.eps_revisions,
        'growth_estimates': lambda: ticker.growth_estimates,
        
        'earnings_dates': lambda: ticker.earnings_dates,
        #msft.isin
        'ISIN' : lambda: pd.DataFrame([{'ISIN': ticker.isin}]),
        #msft.news
        


    }

    # Fetch all data
    for key, fetcher in data_fetchers.items():
        data[key] = safe_fetch(key, fetcher)

    if save_as_excel:
        # Save all data to a single Excel file with multiple sheets
        excel_path = os.path.join(output_folder, f"financials_data.xlsx")
        with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
            for key, df in data.items():
                if not df.empty:
                    try:
                        # Double-check timezone conversion before saving
                        df = make_timezone_naive(df)
                        # Truncate sheet name to Excel's 31 character limit
                        sheet_name = key[:31]
                        df.to_excel(writer, sheet_name=sheet_name, index=True)
                        print(f"Successfully saved sheet: {key}")
                    except Exception as e:
                        print(f"Warning: Failed to save sheet {key}: {str(e)}")
        print(f"All data saved to {excel_path}")
    else:
        # Save each part of the data to separate CSV files
        for key, df in data.items():
            if not df.empty:
                try:
                    # Double-check timezone conversion before saving
                    df = make_timezone_naive(df)
                    csv_path = os.path.join(output_folder, f"financials_{key}.csv")
                    df.to_csv(csv_path, index=True)
                    print(f"{key} data saved to {csv_path}")
                except Exception as e:
                    print(f"Warning: Failed to save {key} to CSV: {str(e)}")

    return data





import pandas as pd
import json
from openai import OpenAI

def generate_financial_risk_analysis(user_prompt= "Please analyze the financial risks.",
                                     api_key= None,
                                     income_statement_file= 'dataspace/financials/financials_income_statement.csv', 
                                     balance_sheet_file= 'dataspace/financials/financials_balance_sheet.csv', 
                                     cash_flow_file = 'dataspace/financials/financials_cashflow.csv',
                                     output_file = 'dataspace/financial_risk_analysis.json', 
                                     model="gpt-4o",
                                     temperature=0.7, 
                                     max_tokens=2048, 
                                     top_p=1):
    # Initialize OpenAI client
    client = OpenAI(api_key=api_key)

    # Load financial reports
    income_statement = pd.read_csv(income_statement_file)
    balance_sheet = pd.read_csv(balance_sheet_file)
    cash_flow = pd.read_csv(cash_flow_file)

    # Convert DataFrames to JSON strings
    income_statement_json = income_statement.to_json(orient='records')
    balance_sheet_json = balance_sheet.to_json(orient='records')
    cash_flow_json = cash_flow.to_json(orient='records')

    # System prompt setup for financial analysis
    system_prompt = {
        "role": "system",
        "content": """
You are a financial analyst AI that conducts comprehensive risk analysis based on provided financial data.

# Instructions

- Analyze the provided income statement, balance sheet, and cash flow statement data.
- Identify potential financial risks, assess their severity and likelihood, and propose mitigation strategies.
- Consider all aspects of the data, including trends over time.

# Output Format

Provide the output as a structured list in JSON format, where each risk includes:

- "section_name": The financial statement section ("Income Statement", "Balance Sheet", "Cash Flow", or "Overall").
- "risk_name": Short title of the risk.
- "description": Detailed description of the risk.
- "severity": Evaluation of severity (e.g., Low, Medium, High).
- "likelihood": Likelihood of the risk occurring (e.g., Low, Medium, High).
- "mitigation": Suggested mitigation strategies.

# Note

- Focus on risks that are evident from the data.
- Ensure the analysis is based on the data provided.
        """
    }

    # Prepare the data to be included in the user prompt
    financial_data_prompt = f"""
The company's financial data is as follows:

- **Income Statement Data**:
{income_statement_json}

- **Balance Sheet Data**:
{balance_sheet_json}

- **Cash Flow Data**:
{cash_flow_json}

{user_prompt}
    """

    user_prompt_data = {
        "role": "user",
        "content": financial_data_prompt
    }

    # Call the OpenAI API
    response = client.chat.completions.create(
        model=model,
        messages=[system_prompt, user_prompt_data],
        temperature=temperature,
        max_tokens=max_tokens,
        top_p=top_p,
        frequency_penalty=0,
        presence_penalty=0
    )

    # Extract JSON content from the response
    try:
        # Access the 'content' attribute directly
        content = response.choices[0].message.content.strip("```json\n").strip("\n```")
        # Parse the content to JSON
        analysis_data = json.loads(content)
    except json.JSONDecodeError:
        print("Failed to decode JSON from the response. Here is the raw response:")
        print(content)
        return

    # Save to JSON file
    with open(output_file, 'w') as f:
        json.dump(analysis_data, f, indent=4)

    print(f"Financial risk analysis saved to {output_file}")



import json
import csv

def json_to_csv(json_file ='dataspace/financial_risk_analysis.json', 
                output_csv_file ='dataspace/financial_risk_analysis.csv' ):
    """
    Converts a JSON file to a CSV file.
    
    Parameters:
    - json_file: Path to the JSON file containing data.
    - output_csv_file: Path to the output CSV file.
    """
    # Load JSON data from the file
    with open(json_file, 'r') as f:
        data = json.load(f)

    # Open the CSV file for writing
    with open(output_csv_file, 'w', newline='') as csvfile:
        # Get the field names from the keys of the first dictionary in the JSON list
        fieldnames = data[0].keys()
        
        # Create the CSV writer object
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
        
        # Write the header
        writer.writeheader()
        
        # Write each dictionary in the JSON data as a row in the CSV
        for row in data:
            writer.writerow(row)
    
    return "Data successfully converted to CSV and saved to {output_csv_file}"

# # Example usage
# json_to_csv('dataspace/financial_risk_analysis.json', 
#             'dataspace/financial_risk_analysis.csv')


import os
import json
import pandas as pd
from typing import Union, Dict, List
from openai import OpenAI
from bs4 import BeautifulSoup

def custom_file_analyzer(
    files: Union[str, List[str]]= "dataspace/dataset1.csv",
    analysis_task: str= "Analyze the content of the file and provide insights.",
    api_key: str = '',
    output_file: str = 'analysis_results.json',
    model: str = 'gpt-4o',
    max_tokens: int = 2000,
    temperature: float = 0.7
) -> Dict:
    """
    Analyzes content of given files based on a user-defined analysis task.

    Parameters:
    - files: A single file path or a list of file paths to analyze.
    - analysis_task: A string describing the analysis to be performed on the content.
    - api_key: OpenAI API key.
    - output_file: Path to save the analysis results.
    - model: OpenAI model to use.
    - max_tokens: Maximum tokens in response.
    - temperature: Response randomness (0-1).

    Returns:
    - Dict containing analysis results.
    """
    # Helper function to check if a module is installed
    def _check_module(module_name: str) -> bool:
        import importlib.util
        return importlib.util.find_spec(module_name) is not None

    # Helper function to extract text from HTML content
    def _extract_text_from_html(html_content: str) -> str:
        soup = BeautifulSoup(html_content, 'html.parser')
        for script in soup(["script", "style"]):
            script.decompose()  # Remove script and style elements
        return soup.get_text(separator='\n', strip=True)

    # Helper function to read file content
    def _read_file_content(file_path: str, installed_modules: Dict[str, bool]) -> str:
        file_extension = os.path.splitext(file_path)[1].lower()

        try:
            if file_extension == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()

            elif file_extension == '.json':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return json.dumps(json.load(f), indent=2)

            elif file_extension == '.html':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return _extract_text_from_html(f.read())

            elif file_extension == '.pdf' and installed_modules.get('pdfplumber'):
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    return '\n'.join(page.extract_text() for page in pdf.pages if page.extract_text())

            elif file_extension == '.docx' and installed_modules.get('docx'):
                import docx
                doc = docx.Document(file_path)
                return '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])

            elif file_extension in ['.csv', '.xlsx'] and installed_modules.get('pandas'):
                if file_extension == '.csv':
                    df = pd.read_csv(file_path)
                else:
                    df = pd.read_excel(file_path)
                return df.to_string()

            else:
                raise ValueError(f"Unsupported file type: {file_extension}")

        except Exception as e:
            raise ValueError(f"Error reading {file_extension} file: {str(e)}")

    # Start of the main function logic
    try:
        # Initialize modules and file types
        installed_modules = {
            module: _check_module(module)
            for module in ['docx', 'pdfplumber', 'pandas', 'openpyxl', 'bs4']
        }

        supported_file_types = ['.txt', '.json', '.html']
        if installed_modules.get('docx'):
            supported_file_types.append('.docx')
        if installed_modules.get('pdfplumber'):
            supported_file_types.append('.pdf')
        if installed_modules.get('pandas'):
            supported_file_types.extend(['.csv', '.xlsx'])

        if not api_key:
            raise ValueError("OpenAI API key is required")

        client = OpenAI(api_key=api_key)
        analysis_results = {}

        # Ensure files is a list
        files = [files] if isinstance(files, str) else files

        # Process each file
        for file_path in files:
            if not os.path.isfile(file_path):
                analysis_results[file_path] = {'error': 'File does not exist'}
                continue

            file_extension = os.path.splitext(file_path)[1].lower()
            if file_extension not in supported_file_types:
                analysis_results[file_path] = {'error': f'Unsupported file type: {file_extension}'}
                continue

            # Read the file content
            content = _read_file_content(file_path, installed_modules)

            # Prepare the message for OpenAI API
            messages = [
                {"role": "system", "content": "You are an expert analyst."},
                {"role": "user", "content": f"Please perform the following analysis task on the content:\n\n{analysis_task}\n\nContent:\n{content}"}
            ]

            # Call the OpenAI API
            response = client.chat.completions.create(
                model=model,
                messages=messages,
                max_tokens=max_tokens,
                temperature=temperature
            )

            # Extract the analysis result
            analysis = response.choices[0].message.content

            # Store the result
            analysis_results[file_path] = {
                'analysis': analysis,
                'file_type': file_extension
            }

        # Save the analysis results to the output file
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(analysis_results, f, indent=4, ensure_ascii=False)

        return analysis_results

    except Exception as e:
        error_result = {'error': str(e)}
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(error_result, f, indent=4, ensure_ascii=False)
        return error_result



import pandas as pd
import yfinance as yf
import os
from datetime import datetime
from typing import Optional, Union, Dict

def __make_timezone_naive(df: pd.DataFrame) -> pd.DataFrame:
    """Helper function to convert DataFrame to timezone-naive"""
    if not isinstance(df, pd.DataFrame):
        if isinstance(df, pd.Series):
            df = df.to_frame()
        else:
            return df
            
    if isinstance(df.index, pd.DatetimeIndex) and df.index.tz is not None:
        df.index = df.index.tz_localize(None)
    
    for col in df.select_dtypes(include=['datetime64[ns, UTC]', 'datetimetz']).columns:
        df[col] = df[col].dt.tz_localize(None)
        
    for col in df.select_dtypes(include=['object']).columns:
        try:
            if pd.api.types.is_datetime64_any_dtype(df[col]):
                df[col] = pd.to_datetime(df[col]).dt.tz_localize(None)
        except:
            continue
            
    return df

def __create_output_dir(output_path: str) -> None:
    """Create output directory if it doesn't exist"""
    if output_path:
        os.makedirs(os.path.dirname(output_path), exist_ok=True)

def __save_dataframe(df: pd.DataFrame, 
                    output_path: str, 
                    sheet_name: str = 'Sheet1',
                    index: bool = True) -> None:
    """
    Save DataFrame to either CSV or Excel
    """
    if not output_path:
        return
        
    __create_output_dir(output_path)
    
    if output_path.endswith('.xlsx'):
        with pd.ExcelWriter(output_path, engine='openpyxl', 
                           mode='a' if os.path.exists(output_path) else 'w') as writer:
            df.to_excel(writer, sheet_name=sheet_name, index=index)
    else:
        df.to_csv(output_path, index=index)

def __get_safe_filename(base: str, name: str, ext: str) -> str:
    """Create safe filename by removing invalid characters"""
    invalid_chars = '<>:"/\\|?*'
    for char in invalid_chars:
        name = name.replace(char, '_')
    return f"{base}_{name}{ext}"

def get_basic_info(ticker_symbol: str = 'MSFT', 
                  output_path: str = 'dataspace/info.csv') -> pd.DataFrame:
    """Get basic information about the stock"""
    ticker = yf.Ticker(ticker_symbol)
    info_df = pd.DataFrame.from_dict(ticker.info, orient='index', columns=[ticker_symbol])
    info_df = __make_timezone_naive(info_df)
    
    if output_path:
        __save_dataframe(info_df, output_path)
    
    return info_df

def get_market_data(ticker_symbol: str = 'MSFT', 
                   period: str = '1y', 
                   output_path: str = 'dataspace/market_data.csv') -> pd.DataFrame:
    """Get historical market data"""
    ticker = yf.Ticker(ticker_symbol)
    market_data = ticker.history(period=period)
    market_data = __make_timezone_naive(market_data)
    
    if output_path:
        __save_dataframe(market_data, output_path)
    
    return market_data

def get_financial_statements(ticker_symbol: str = 'MSFT', 
                           statement_type: str = 'all', 
                           frequency: str = 'annual', 
                           output_path: str = 'dataspace/financial_statements.xlsx') -> Dict[str, pd.DataFrame]:
    """Get financial statements"""
    ticker = yf.Ticker(ticker_symbol)
    statements = {}
    
    statement_mapping = {
        'income': (ticker.income_stmt, ticker.quarterly_income_stmt),
        'balance': (ticker.balance_sheet, ticker.quarterly_balance_sheet),
        'cash': (ticker.cashflow, ticker.quarterly_cashflow)
    }
    
    statements_to_fetch = (
        statement_mapping.items() if statement_type == 'all' 
        else [(statement_type, statement_mapping[statement_type])]
    )
    
    for name, (annual, quarterly) in statements_to_fetch:
        df = __make_timezone_naive(
            (annual if frequency == 'annual' else quarterly).transpose()
        )
        statements[name] = df
        
        if output_path:
            if output_path.endswith('.xlsx'):
                __save_dataframe(df, output_path, sheet_name=f"{frequency}_{name}")
            else:
                base, ext = os.path.splitext(output_path)
                out_path = __get_safe_filename(base, f"{frequency}_{name}", ext)
                __save_dataframe(df, out_path)
    
    return statements

def get_ownership_data(ticker_symbol: str = 'MSFT', 
                      data_type: str = 'all', 
                      output_path: str = 'dataspace/ownership.xlsx') -> Dict[str, pd.DataFrame]:
    """Get ownership and insider data"""
    ticker = yf.Ticker(ticker_symbol)
    ownership = {}
    
    data_mapping = {
        'major': {'major_holders': ticker.major_holders},
        'institutional': {'institutional_holders': ticker.institutional_holders},
        'mutual': {'mutualfund_holders': ticker.mutualfund_holders},
        'insider': {
            'insider_transactions': ticker.insider_transactions,
            'insider_purchases': ticker.insider_purchases,
            'insider_roster': ticker.insider_roster_holders
        }
    }
    
    types_to_fetch = (
        data_mapping.items() if data_type == 'all' 
        else [(data_type, data_mapping[data_type])]
    )
    
    for _, data_dict in types_to_fetch:
        for name, data in data_dict.items():
            df = __make_timezone_naive(data)
            ownership[name] = df
            
            if output_path:
                if output_path.endswith('.xlsx'):
                    __save_dataframe(df, output_path, sheet_name=name)
                else:
                    base, ext = os.path.splitext(output_path)
                    out_path = __get_safe_filename(base, name, ext)
                    __save_dataframe(df, out_path)
    
    return ownership

def get_analyst_data(ticker_symbol: str = 'MSFT', 
                    data_type: str = 'all', 
                    output_path: str = 'dataspace/analyst.xlsx') -> Dict[str, pd.DataFrame]:
    """Get analyst recommendations and estimates"""
    ticker = yf.Ticker(ticker_symbol)
    analysis = {}
    
    data_mapping = {
        'recommendations': {
            'recommendations': ticker.recommendations,
            'recommendations_summary': ticker.recommendations_summary,
            'upgrades_downgrades': ticker.upgrades_downgrades
        },
        'price_targets': {
            'price_targets': ticker.analyst_price_targets
        },
        'estimates': {
            'earnings_estimate': ticker.earnings_estimate,
            'revenue_estimate': ticker.revenue_estimate,
            'earnings_history': ticker.earnings_history,
            'eps_trend': ticker.eps_trend,
            'eps_revisions': ticker.eps_revisions,
            'growth_estimates': ticker.growth_estimates
        }
    }
    
    types_to_fetch = (
        data_mapping.items() if data_type == 'all' 
        else [(data_type, data_mapping[data_type])]
    )
    
    for _, data_dict in types_to_fetch:
        for name, data in data_dict.items():
            df = __make_timezone_naive(data)
            analysis[name] = df
            
            if output_path:
                if output_path.endswith('.xlsx'):
                    __save_dataframe(df, output_path, sheet_name=name)
                else:
                    base, ext = os.path.splitext(output_path)
                    out_path = __get_safe_filename(base, name, ext)
                    __save_dataframe(df, out_path)
    
    return analysis

def get_corporate_actions(ticker_symbol: str = 'MSFT', 
                         data_type: str = 'all', 
                         output_path: str = 'dataspace/corporate_actions.xlsx') -> Dict[str, pd.DataFrame]:
    """Get corporate actions data"""
    ticker = yf.Ticker(ticker_symbol)
    actions = {}
    
    data_mapping = {
        'dividends': {
            'dividends': ticker.dividends.to_frame() if not ticker.dividends.empty else pd.DataFrame()
        },
        'splits': {
            'splits': ticker.splits.to_frame() if not ticker.splits.empty else pd.DataFrame()
        },
        'capital_gains': {
            'capital_gains': ticker.capital_gains.to_frame() if not ticker.capital_gains.empty else pd.DataFrame()
        }
    }
    
    types_to_fetch = (
        data_mapping.items() if data_type == 'all' 
        else [(data_type, data_mapping[data_type])]
    )
    
    for _, data_dict in types_to_fetch:
        for name, data in data_dict.items():
            df = __make_timezone_naive(data)
            actions[name] = df
            
            if output_path:
                if output_path.endswith('.xlsx'):
                    __save_dataframe(df, output_path, sheet_name=name)
                else:
                    base, ext = os.path.splitext(output_path)
                    out_path = __get_safe_filename(base, name, ext)
                    __save_dataframe(df, out_path)
    
    return actions

def get_news_data(ticker_symbol: str = 'MSFT', 
                  output_path: str = 'dataspace/news.csv') -> pd.DataFrame:
    """Get news related to the stock"""
    ticker = yf.Ticker(ticker_symbol)
    news_df = pd.DataFrame(ticker.news)
    news_df = __make_timezone_naive(news_df)
    
    if output_path:
        __save_dataframe(news_df, output_path)
    
    return news_df

def get_calendar_data(ticker_symbol: str = 'MSFT', 
                     output_path: str = 'dataspace/calendar.csv') -> pd.DataFrame:
    """Get calendar events"""
    ticker = yf.Ticker(ticker_symbol)
    calendar_df = pd.DataFrame(ticker.calendar)
    calendar_df = __make_timezone_naive(calendar_df)
    
    if output_path:
        __save_dataframe(calendar_df, output_path)
    
    return calendar_df

def get_sec_filings(ticker_symbol: str = 'MSFT', 
                    output_path: str = 'dataspace/sec_filings.csv') -> pd.DataFrame:
    """Get SEC filings"""
    ticker = yf.Ticker(ticker_symbol)
    filings_df = pd.DataFrame(ticker.sec_filings)
    filings_df = __make_timezone_naive(filings_df)
    
    if output_path:
        __save_dataframe(filings_df, output_path)
    
    return filings_df

